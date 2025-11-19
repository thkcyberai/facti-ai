from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_heading('FACTI.AI KYC MODEL', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_heading('Technical Development Roadmap', level=2)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('Pure Technical Implementation Guide')
doc.add_paragraph('November 2025')
doc.add_paragraph()

# Overview
doc.add_heading('OVERVIEW: KYC MODEL ARCHITECTURE', 1)
p = doc.add_paragraph('The KYC Model is an ')
p.add_run('ENSEMBLE').bold = True
p.add_run(' of 4 components working together:')

doc.add_paragraph('1. Document Authentication → Is ID real?', style='List Bullet')
doc.add_paragraph('2. Video Liveness Detection → Is person real?', style='List Bullet')
doc.add_paragraph('3. Face Matching → Does ID photo = video face?', style='List Bullet')
doc.add_paragraph('4. Fraud Scoring → Suspicious patterns?', style='List Bullet')
doc.add_paragraph('Final Verdict: PASS / FAIL / REVIEW', style='List Bullet')

doc.add_page_break()

# Component 1
doc.add_heading('COMPONENT 1: DOCUMENT AUTHENTICATION', 1)
doc.add_paragraph('Goal: Detect fake/forged ID documents (95%+ accuracy)')

doc.add_heading('What to Detect:', 2)
doc.add_paragraph('Photo substitution on IDs', style='List Bullet')
doc.add_paragraph('Template forgeries', style='List Bullet')
doc.add_paragraph('Security feature tampering', style='List Bullet')
doc.add_paragraph('Text alterations', style='List Bullet')

doc.add_heading('Training Cost: $27', 3)
doc.add_heading('Time: 3-4 days', 3)

doc.add_page_break()

# Component 2
doc.add_heading('COMPONENT 2: VIDEO LIVENESS DETECTION', 1)
doc.add_paragraph('Goal: Prove real person (99%+ accuracy)')

doc.add_heading('What to Detect:', 2)
doc.add_paragraph('Static photo attacks', style='List Bullet')
doc.add_paragraph('Video replay attacks', style='List Bullet')
doc.add_paragraph('Deepfake videos', style='List Bullet')

doc.add_heading('Training Cost: $54', 3)
doc.add_heading('Time: 5-7 days', 3)

doc.add_page_break()

# Component 3
doc.add_heading('COMPONENT 3: FACE MATCHING', 1)
doc.add_paragraph('Goal: Match ID photo to video face (98%+ accuracy)')
p = doc.add_paragraph()
p.add_run('Cost: $0 (pre-trained model)').bold = True

doc.add_page_break()

# Component 4
doc.add_heading('COMPONENT 4: FRAUD SCORING', 1)
doc.add_paragraph('Goal: Detect suspicious patterns')
p = doc.add_paragraph()
p.add_run('Cost: $0 (CPU training)').bold = True

doc.add_page_break()

# Summary
doc.add_heading('PROJECT SUMMARY', 1)
p = doc.add_paragraph()
p.add_run('TOTAL COST: $81').bold = True
doc.add_paragraph('Time: 6-8 weeks')

# Save
doc.save('KYC_Model_Technical_Roadmap.docx')
print("Document created successfully!")
